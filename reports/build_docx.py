"""Generate a concise analyst-style Word note for UBER."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK      = RGBColor(0x1f, 0x23, 0x2e)
MUTE     = RGBColor(0x60, 0x68, 0x78)
ACCENT   = RGBColor(0x6d, 0x28, 0xd9)
POS      = RGBColor(0x15, 0x80, 0x3d)
NEG      = RGBColor(0xb9, 0x1c, 0x1c)
WARN     = RGBColor(0xb4, 0x53, 0x09)
LIGHT_BG = "F4F5F8"
ACCENT_BG = "EDE7FE"

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

# Base style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.font.color.rgb = INK


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_run(p, text, bold=False, color=None, size=None, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return r


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, bold=True, size=18, color=INK)


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text, bold=True, size=12, color=ACCENT)


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    add_run(p, text, bold=True, size=10.5, color=INK)


def bullet(text_runs, indent=0):
    """text_runs: list of (text, bold, color) tuples."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
    for t in text_runs:
        if isinstance(t, str):
            add_run(p, t)
        else:
            text, bold, color = (t + (None,) * 3)[:3]
            add_run(p, text, bold=bold or False, color=color)


def kv_table(rows, col_widths=None, header=None, header_bg=ACCENT_BG):
    """rows: list of tuples. Optional header tuple."""
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows) + (1 if header else 0), cols=n_cols)
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for c in table.columns[i].cells:
                c.width = Cm(w)
    start = 0
    if header:
        for i, h in enumerate(header):
            c = table.rows[0].cells[i]
            c.text = ''
            p = c.paragraphs[0]
            add_run(p, h, bold=True, size=9, color=INK)
            set_cell_bg(c, header_bg)
        start = 1
    for r_idx, row in enumerate(rows):
        for i, val in enumerate(row):
            c = table.rows[r_idx + start].cells[i]
            c.text = ''
            p = c.paragraphs[0]
            if isinstance(val, tuple):
                text, color = val
                add_run(p, text, size=9.5, color=color)
            else:
                add_run(p, str(val), size=9.5)
            if r_idx % 2 == 1:
                set_cell_bg(c, LIGHT_BG)
    # borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'D5D8DF')
        borders.append(b)
    tblPr.append(borders)
    return table


# ====== HEADER BLOCK ======
h1("Uber Technologies, Inc.   |   NYSE: UBER")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
add_run(p, "Equity Analyst Note  ·  ", color=MUTE, size=9)
add_run(p, "May 14, 2026", bold=True, size=9, color=MUTE)
add_run(p, "  ·  Framework: V. Khandelwal 15-Q template", color=MUTE, size=9)

# Recommendation banner
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
add_run(p, "Rating:  ", size=10, color=MUTE)
add_run(p, "HOLD", bold=True, size=12, color=WARN)
add_run(p, "    Buy < $65  ·  Load < $55  ·  Trim > $120", size=10, color=MUTE)

# Snapshot table
kv_table(
    rows=[
        ("Price", "$74.69", "Market cap", "$152.0B"),
        ("52-wk Δ", ("-18.6%", NEG), "Enterprise value", "$154.5B"),
        ("TTM Revenue", "$53.7B", "TTM FCF", ("$9.8B", POS)),
        ("EV / FCF", ("15.8×", POS), "FCF yield", ("6.45%", POS)),
        ("ROIC / WACC", ("19.3% / 10.0%", POS), "Net debt", "$2.5B"),
        ("Analyst PT (cons.)", ("$107.79 (+44%)", POS), "Rating", "Strong Buy (34)"),
    ],
    col_widths=[3.6, 4.0, 3.8, 4.0],
)

# ====== KEY TAKEAWAYS ======
h2("Key Takeaways")
bullet([("Thesis. ", True), "A good business at a fair price — not a great business at a great price. Asset-light marketplace generating $10B+ FCF on $54B revenue, compounding ~25% CAGR over 4 years."])
bullet([("Catalyst-light here. ", True), "Reverse-DCF implies 10–12% FCF CAGR for 10y; my base case 15–18% near-term then 10–12%. Market is pricing roughly what I'd model — no asymmetry today."])
bullet([("Main debate. ", True), "AV economics in 5–10 years. Uber positioned on both sides (Waymo/Wayve/Aurora partner + largest demand network). Take-rate compression is the bear case; advertising mix-up is the bull case."])
bullet([("Position. ", True), "Starter 2–3% at current; scale 5–6% on AV-fear sell-off below $65; load below $55."])

# ====== BUSINESS ======
h2("Section 1 · Business")

h3("What it does, how it earns")
bullet([("Platform matching: ", True), "riders↔drivers (Mobility), eaters↔restaurants↔couriers (Delivery), shippers↔truckers (Freight). Asset-light; ~25–30¢ take of every transacted dollar."])

kv_table(
    header=("Segment", "% Revenue", "Take Rate", "Seg EBITDA mgn"),
    rows=[
        ("Mobility", "~57%", "~30%", "~7–8%"),
        ("Delivery", "~37%", "~19–20%", "~3.5–4%"),
        ("Freight", "~6%", "~100% gross", "~0%"),
        ("Ads (embedded)", "~3%", "—", "Very high incremental"),
    ],
    col_widths=[3.5, 3.8, 4.0, 4.1],
)

h3("What could kill it (10y view)")
kv_table(
    header=("Risk", "Prob.", "Severity", "Mitigant"),
    rows=[
        ("AVs bypass demand layer (Tesla/Waymo direct-to-consumer)", "Med", ("Severe", NEG), "Consumer app switching cost; demand-pattern data"),
        ("Driver reclassification (PRO Act, EU PWD)", ("Med-High", WARN), ("High", NEG), "Survived AB5; can re-price"),
        ("Insurance cost spiral (CA/NJ)", "Med", "Med", "Consumer-paid surcharges piloting"),
        ("Recession crushes Delivery", "Med", "Med", "Now profitable, not a cash sink"),
    ],
    col_widths=[6.5, 2.0, 2.2, 4.7],
)

h3("Honest understanding (~80%)")
bullet([("Understood: ", True), "two-sided network effects; why Mobility > Delivery economics; asset-light FCF compounding."])
bullet([("Not fully: ", True), "gross-vs-net revenue accounting by jurisdiction; AV-era unit economics; insurance reserve actuarial volatility."])

# ====== MANAGEMENT ======
h2("Section 2 · Management")

h3("Quality")
bullet([("Dara Khosrowshahi ", True), "(CEO since Aug-2017): clean tenure, turnaround executed. From Expedia (2005–17). New CFO ", ("Prashanth Mahendra-Rajah ", True), "(ex-Analog Devices, 2024) signals shift from growth-CFO to discipline-CFO."])

h3("Capital allocation track record")
kv_table(
    header=("Era", "Move", "Grade"),
    rows=[
        ("2014–18 · Kalanick", "Burned ~$10B on China, SE Asia, ATG", ("Disaster (D)", NEG)),
        ("2018–20 · Dara P1", "Sold China→Didi, SE Asia→Grab, ATG→Aurora for equity (~$8B LT inv. today)", ("Excellent (A)", POS)),
        ("2021–23 · Dara P2", "'Sea change' memo; cost discipline; no speculative M&A", ("Strong (A-)", POS)),
        ("2024–26 · Dara P3", "Initiated buybacks ($1.25B FY24 → $6.5B FY25); Foodpanda Taiwan blocked", ("Good (B+)", WARN)),
    ],
    col_widths=[3.4, 9.2, 2.8],
)

h3("Minority shareholders & honesty")
bullet([("+ ", True, POS), "Killed dual-class supervote (2023); buybacks scaled to $6.5B FY25; segment-level economics disclosed in plain English since 2022."])
bullet([("+ ", True, POS), "Communicated honestly in good years and bad — FY24 letter led with the $5.8B deferred-tax valuation allowance benefit rather than hiding it."])
bullet([("− ", True, NEG), "SBC $1.86B/yr (~19% of FCF) reduces real shareholder yield from 5% → ~3%."])
bullet([("− ", True, NEG), "Insider ownership only 0.17%; no founder remains. Travis Kalanick exited entire 8.6% by 2020. Skin in game is weak."])

# ====== PRICE ======
h2("Section 3 · Price")

h3("Valuation snapshot")
kv_table(
    header=("Multiple", "Value", "Read"),
    rows=[
        ("Trailing P/E", "18.5×", "Distorted by tax assets"),
        ("Fwd P/E (2026 cons.)", "21.7×", "Reset year (no tax benefit)"),
        ("P/E (2027 cons.)", ("16.9×", POS), "Cleaner lens"),
        ("EV / EBITDA", "22.0×", "Optically high"),
        ("EV / FCF", ("15.8×", POS), "Right lens — asset-light"),
        ("PEG", ("0.64", POS), "Cheap if growth holds"),
        ("FCF yield + growth", ("6.45% + 12–15%", POS), "Mid-teens fwd return potential"),
    ],
    col_widths=[4.0, 4.0, 7.4],
)

h3("Reverse-DCF & private-market value")
bullet([("Reverse-DCF: ", True), "today's price requires ~10–12% FCF CAGR for 10y, then 3% terminal."])
bullet([("My base case: ", True), "15–18% near-term, 10–12% mid-term FCF growth. Market is pricing what I'd model — fair, not cheap."])
bullet([("Private-market value: ", True), "Normalized FCF ~$11B ÷ (10%−3%) = EV $157B → equity ~$74/share. Today ≈ rational private buyer."])

h3("Action plan")
kv_table(
    header=("Trigger", "Action"),
    rows=[
        ("Price > $120", ("Trim 50% (≈ 25× fwd EV/FCF; prices in flawless AV transition)", NEG)),
        ("FCF margin contracts 2 consecutive Qs (no one-off)", ("Exit — AV take-rate compression real", NEG)),
        (">$10B acquisition / re-entry China / own AV fleet", ("Exit — capital allocation thesis broken", NEG)),
        ("Price < $65 (~13× EV/FCF)", ("Add — 13% margin of safety to private-market value", POS)),
        ("Price < $55 (~11× EV/FCF)", ("Load — back up the truck", POS)),
        ("-40% on macro/sentiment only", ("Buy aggressively; pre-commitment beats analysis", POS)),
    ],
    col_widths=[7.0, 8.4],
)

# ====== SCORECARD ======
h2("Scorecard")
kv_table(
    header=("Dimension", "Score (/5)", "Note"),
    rows=[
        ("Business quality", ("4.0", POS), "Real moat, $10B FCF, asset-light"),
        ("Business durability (10y)", ("3.0", WARN), "AV transition is a real unknown"),
        ("Management quality", ("4.0", POS), "Honest, disciplined post-2018"),
        ("Skin in game", ("2.0", NEG), "0.17% insider ownership"),
        ("Capital allocation (track record)", ("3.5", WARN), "Improving; not yet proven in maturity"),
        ("Price today", ("3.0", WARN), "Fair, not cheap"),
        ("Price if -40%", ("5.0", POS), "Strong buy zone"),
        ("Upside / downside asymmetry", ("3.5", WARN), "Skewed slightly favorable"),
    ],
    col_widths=[6.0, 2.5, 6.9],
)

# Footer
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
add_run(p, "Educational analysis · Not investment advice · Data: company filings, stockanalysis.com (May 14, 2026)", italic=True, size=8, color=MUTE)

out = r"C:\Users\hariragu\playground\analyst\UBER_Analysis.docx"
doc.save(out)
print(f"Saved: {out}")
